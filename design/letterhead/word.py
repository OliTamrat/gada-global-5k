"""
Editable Word versions of the letterhead — the file the organizers type into.

    pip install python-docx                 # build-time only, not a site dependency
    python3 design/letterhead/build.py      # sponsors.ts -> HTML sheets
    node    design/letterhead/render.mjs    # -> 300dpi PNGs
    python3 design/letterhead/word.py       # -> .docx

Two files come out:

    gada-5k-letterhead-blank.docx   an empty sheet to write any letter on
    gada-5k-letterhead-letter.docx  the model sponsorship request, editable

The artwork is not pasted into the body. It is one full-page image placed in the
*page header*, anchored to the page and sitting behind the text. That matters
for a document other people will edit: header art cannot be selected, dragged or
deleted by someone typing in the body, it repeats by itself if the letter runs
to a second page, and the margins are taken from the design, so typed text lands
in the same column as the printed sample.

python-docx has no API for a behind-text image, so the picture goes in as an
inline drawing and the ``wp:inline`` element is then swapped for a ``wp:anchor``.
Everything fiddly about the package — content types, relationships, the media
part — stays with the library, and only that one element is hand-written.

The letter wording is imported from ``build.py`` rather than repeated here, so
the Word file and the printed sheet cannot drift apart.

Note the module is *word.py*, not docx.py: a file named docx.py next to this one
shadows the python-docx library and the import silently resolves to the wrong
module.
"""
import copy
import pathlib

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu, Inches, Pt, RGBColor

from build import LETTER_BLOCKS

OUT = pathlib.Path(__file__).resolve().parent

# ── Geometry, measured on the rendered sheet ──────────────────────────────
# The HTML sheet is 850x1100 px for 8.5x11 in, so 100 px is exactly one inch.
# The text column runs from y=173 to the footer rule at y=1010.9, x=68 to 782.
PAGE_W, PAGE_H = 8.5, 11.0
MARGIN_X = 0.68
MARGIN_TOP = 1.73
MARGIN_BOTTOM = 0.95      # the footer rule is at 0.891 in; the rest is clearance

INK = RGBColor(0x14, 0x12, 0x10)
MUTED = RGBColor(0x5C, 0x56, 0x4C)
GOLD = "F5C842"

# Georgia rather than the IBM Plex Serif of the printed sheet: it ships with
# Windows and macOS, so the file looks right on a machine that has installed
# nothing. The Plex fonts sit beside this script for an exact match — install
# them, then select all and switch the font.
BODY_FONT = "Georgia"
MONO_FONT = "Consolas"

BODY_PT = Pt(10)
SMALL_PT = Pt(9.5)
LEVELS_PT = Pt(9)

# Everything that may legally follow w:pBdr inside w:pPr. python-docx uses this
# to insert the border in its schema-mandated position; order is not optional in
# OOXML and Word rejects a file that gets it wrong.
AFTER_PBDR = (
    "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)


def add_para(doc, text="", *, font=BODY_FONT, size=BODY_PT, color=INK,
             bold=False, after=9, line=1.5, underline=False, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.left_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = size
        r.font.color.rgb = color
        r.font.bold = bold
        r.font.underline = underline
        # Word consults rFonts/@w:cs for anything it treats as complex script;
        # without it a run can silently fall back to the theme font.
        r._r.rPr.rFonts.set(qn("w:cs"), font)
    return p


def add_left_rule(paragraph, color=GOLD, width_eighths=18, gap=10):
    """The gold rule beside the levels line."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{width_eighths}" w:space="{gap}" '
        f'w:color="{color}"/></w:pBdr>')
    pPr.insert_element_before(pBdr, *AFTER_PBDR)


def background_in_header(section, image: pathlib.Path):
    """Put the sheet art behind the text, locked, repeating on every page."""
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run()
    run.add_picture(str(image), width=Inches(PAGE_W), height=Inches(PAGE_H))

    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    graphic = copy.deepcopy(inline.find(qn("a:graphic")))
    cx, cy = Emu(Inches(PAGE_W)), Emu(Inches(PAGE_H))

    anchor = parse_xml(
        f'<wp:anchor {nsdecls("wp", "a", "pic", "r")} '
        f'distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        f'relativeHeight="0" behindDoc="1" locked="1" layoutInCell="1" '
        f'allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="1" name="Letterhead" descr="Gada Global Inc. letterhead"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks noChangeAspect="1" noMove="1" noResize="1" noSelect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'</wp:anchor>')
    anchor.append(graphic)          # graphic is the last child of wp:anchor
    drawing.remove(inline)
    drawing.append(anchor)


def new_sheet(background: pathlib.Path, title: str) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_PT
    normal.font.color.rgb = INK
    normal.element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(PAGE_W), Inches(PAGE_H)
    sec.left_margin = sec.right_margin = Inches(MARGIN_X)
    sec.top_margin = Inches(MARGIN_TOP)
    sec.bottom_margin = Inches(MARGIN_BOTTOM)
    sec.header_distance = Emu(0)
    sec.footer_distance = Emu(0)

    background_in_header(sec, background)

    doc.core_properties.title = title
    doc.core_properties.author = "Gada Global Inc."
    return doc


def fill_letter(doc):
    for kind, value in LETTER_BLOCKS:
        if kind == "meta":
            for line in value:
                add_para(doc, line, size=SMALL_PT, color=MUTED, after=0, line=1.35)
            add_para(doc, after=10)
        elif kind == "subject":
            add_para(doc, value, bold=True, after=10)
        elif kind == "p":
            add_para(doc, value)
        elif kind == "levels":
            p = add_para(doc, value, font=MONO_FONT, size=LEVELS_PT,
                         after=11, line=1.4, indent=0.14)
            add_left_rule(p)
        elif kind == "sign":
            add_para(doc, value[0], after=3)
            # A ruled line to sign above: underlined spaces, so retyping the
            # name underneath cannot drag the rule with it.
            add_para(doc, " " * 44, underline=True, color=MUTED, after=6, line=1.2)
            for line in value[1:]:
                add_para(doc, line, size=SMALL_PT, after=0, line=1.35)


def main() -> None:
    background = OUT / "gada-5k-letterhead-blank.png"
    if not background.exists():
        raise SystemExit(
            f"{background.name} is missing — run build.py, then render.mjs, first")

    blank = new_sheet(background, "Gada Global Inc. letterhead")
    # A genuinely empty sheet: one paragraph so the cursor opens in the right
    # place. Placeholder text here would only be something to delete first.
    add_para(blank)
    blank.save(OUT / "gada-5k-letterhead-blank.docx")

    letter = new_sheet(background, "Gada Global 5K Run — sponsorship request")
    fill_letter(letter)
    letter.save(OUT / "gada-5k-letterhead-letter.docx")

    for name in ("gada-5k-letterhead-blank.docx", "gada-5k-letterhead-letter.docx"):
        print(f"{name}  {(OUT / name).stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
